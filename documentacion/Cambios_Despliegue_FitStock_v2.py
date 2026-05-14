from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MEMORIA DE CAMBIOS\nDespliegue FitStock')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Migración de Angular SSR a SPA Estático\nConfiguración de API en Producción')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_page_break()

doc.add_heading('Índice', level=1)
items = [
    '1. Resumen de los cambios realizados',
    '2. Cambio 1: URLs de los servicios Angular (localhost -> producción)',
    '3. Cambio 2: Desactivación de SSR (Server-Side Rendering)',
    '4. Cambio 3: Actualización del router.php',
    '5. Cambio 4: Actualización de api/index.php (routing desde subcarpeta)',
    '6. Cambio 5: Corrección de URLs de imágenes de productos',
    '7. Cambio 6: Corrección de ruta de subida de imágenes',
    '8. Cambio 7: Build de producción sin SSR',
    '9. Cambio 8: Añadir RewriteBase al .htaccess (error al refrescar)',
    '10. Cambio 9: Credenciales de base de datos (conexion.php)',
    '11. Cambio 10: Cabeceras CORS en router.php',
    '12. Cambio 11: CSS inline en documentación (subcarpeta)',
    '13. Cambio 12: Permisos de archivos en el servidor',
    '14. Registro de errores durante el despliegue',
    '15. Archivos a subir por FTP',
    '16. Estructura final del servidor',
    '17. Verificación post-despliegue',
]
for item in items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

doc.add_heading('1. Resumen de los cambios realizados', level=1)
doc.add_paragraph(
    'Se han realizado los siguientes cambios en el proyecto FitStock para '
    'desplegarlo correctamente en el hosting de Arsys (plan compartido Linux), '
    'con la aplicación alojada en una subcarpeta /API/:'
)

cambios = [
    ('Actualización de URLs de API',
     'Todos los servicios Angular apuntaban a http://localhost:8000/api '
     '(entorno de desarrollo). Se ha cambiado a https://chomsky.es/API/api '
     '(entorno de producción con subcarpeta).'),
    ('Desactivación de SSR',
     'La aplicación usaba Angular SSR, que requiere Node.js en el servidor. '
     'Se ha compilado como SPA estático con --no-server.'),
    ('Actualización del router.php',
     'Modificado para que sirva index.html de Angular en rutas no-API '
     'y para que reconozca peticiones /API/api (case-insensitive).'),
    ('Actualización de api/index.php',
     'Añadido case "API" y normalización de rutas para funcionar desde '
     'la subcarpeta /API/.'),
    ('Corrección de URLs de imágenes',
     'Cambiado de ruta absoluta /images/productos/ a relativa '
     'images/productos/ para que funcione con el base-href /API/.'),
    ('Corrección de ruta de subida de imágenes',
     'Actualizada la ruta donde se guardan las imágenes al subirlas '
     'desde la app para que apunte a la carpeta correcta del servidor.'),
    ('Actualización de la base de datos',
     'Añadidos datos de ejemplo (materiales, productos, incidencias) '
     'al archivo databasearsys.sql.'),
]
for titulo, desc in cambios:
    p = doc.add_paragraph()
    run = p.add_run(f'{titulo}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()
doc.add_heading('2. Cambio 1: URLs de los servicios Angular', level=1)
doc.add_paragraph(
    'Todos los servicios de Angular tenían la URL http://localhost:8000/api '
    '(entorno de desarrollo). Se ha cambiado a https://chomsky.es/API/api '
    'para que apunten al servidor de producción, teniendo en cuenta que '
    'la aplicación está dentro de la subcarpeta /API/.'
)

doc.add_heading('Archivos modificados (7 servicios):', level=2)

servicios = [
    ('compras.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('incidencias.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('materiales.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('prestamos.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('productos.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('resumen.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('usuario.ts', 'private API_URL = \'http://localhost:8000/api\''),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Archivo (src/app/services/)'
hdr[1].text = 'Cambio realizado'
for archivo, linea in servicios:
    row = table.add_row().cells
    row[0].text = archivo
    row[1].text = 'http://localhost:8000/api → https://chomsky.es/API/api'

doc.add_page_break()
doc.add_heading('3. Cambio 2: Desactivación de SSR', level=1)
doc.add_paragraph(
    'La aplicación Angular estaba configurada con SSR. Como el hosting '
    'compartido de Arsys no ejecuta Node.js, se compila como SPA estático.'
)
doc.add_paragraph(
    'Comando utilizado:\n\n  npm run build -- --no-server --base-href /API/\n\n'
    'La flag --no-server evita generar archivos de servidor Node.js.\n'
    'La flag --base-href /API/ asegura que las rutas absolutas de Angular '
    'funcionen desde la subcarpeta /API/.'
)

doc.add_page_break()
doc.add_heading('4. Cambio 3: Actualización del router.php', level=1)
doc.add_paragraph(
    'El router.php es el punto de entrada de todas las peticiones HTTP. '
    'Se han realizado dos cambios respecto a la versión original:'
)

doc.add_heading('4.1. Regex para detectar /api (case-insensitive)', level=2)
doc.add_paragraph(
    'Línea 31. Antes:\n'
    '  preg_match(\'#^/api#\', $uri)\n\n'
    'Ahora:\n'
    '  preg_match(\'#/api(/|$)#i\', $uri)\n\n'
    'Esto permite detectar /api tanto en rutas como /api/productos como '
    'en /API/api/productos (case-insensitive).'
)

doc.add_heading('4.2. Fallback para Angular SPA', level=2)
doc.add_paragraph(
    'Líneas 36-43. Se añadió la sección "Fallback: servir index.html de '
    'Angular". Antes de mostrar la documentación, busca index.html de '
    'Angular en la raíz y lo sirve. Esto permite que rutas como /login, '
    '/admin/inventario, etc. funcionen con el enrutamiento interno de Angular.'
)

doc.add_page_break()
doc.add_heading('5. Cambio 4: Actualización de api/index.php', level=1)
doc.add_paragraph(
    'El api/index.php es el enrutador interno de la API. Se han realizado '
    'dos cambios para que funcione desde la subcarpeta /API/:'
)

doc.add_heading('5.1. Aceptar "API" como action válido', level=2)
doc.add_paragraph(
    'Línea 56. Se añadió:\n'
    '  case \'API\':\n\n'
    'Justo debajo de case \'api\':. Cuando la app está en /API/, '
    'el primer segmento de la ruta es "API" (mayúsculas), no "api".'
)

doc.add_heading('5.2. Normalización de rutas', level=2)
doc.add_paragraph(
    'Líneas 34-36. Se añadió lógica para detectar y eliminar el prefijo '
    '"API" del array de rutas:\n\n'
    '  if (strtoupper($path[0] ?? \'\') === \'API\' && strtolower($path[1] ?? \'\') === \'api\') {\n'
    '      array_shift($path);\n'
    '  }\n\n'
    'Esto normaliza rutas como /API/api/productos a /api/productos, '
    'dejando el array en el formato que espera el resto del código '
    '(path[0]=api, path[1]=recurso, path[2]=ID).'
)

doc.add_page_break()
doc.add_heading('6. Cambio 5: URLs de imágenes de productos', level=1)
doc.add_paragraph(
    'En el componente producto-list.component.ts se corrigió la función '
    'getImagenUrl() que construye la URL de las imágenes de productos.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('src/app/components/producto-list/producto-list.ts, línea 226')

doc.add_heading('Antes:', level=2)
p = doc.add_paragraph()
run = p.add_run('return \'/images/productos/\' + encodeURIComponent(nombre) + \'.jpg\';')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Después:', level=2)
p = doc.add_paragraph()
run = p.add_run('return \'images/productos/\' + encodeURIComponent(nombre) + \'.jpg\';')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph(
    'La diferencia es la barra inicial (/). Con la barra, la URL es '
    'absoluta desde la raíz del dominio (/images/productos/...), pero '
    'la app está en /API/, por lo que no encontraba las imágenes. '
    'Sin la barra, la URL es relativa al base-href /API/, resolviéndose '
    'como /API/images/productos/...'
)

doc.add_page_break()
doc.add_heading('7. Cambio 6: Ruta de subida de imágenes', level=1)
doc.add_paragraph(
    'En api/index.php se corrigió la ruta donde se guardan las imágenes '
    'al subirlas desde la aplicación.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('FitStock-API/api/index.php, línea 386')

doc.add_heading('Antes:', level=2)
p = doc.add_paragraph()
run = p.add_run('$uploadDir = __DIR__ . \'/../../FitStock-APP/public/images/productos/\';')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Después:', level=2)
p = doc.add_paragraph()
run = p.add_run('$uploadDir = __DIR__ . \'/../images/productos/\';')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph(
    'La ruta anterior estaba diseñada para el entorno de desarrollo local '
    '(Docker). En el servidor de producción, las imágenes se almacenan en '
    'la carpeta /html/API/images/productos/, que es donde las busca el '
    'frontend Angular.'
)

doc.add_page_break()
doc.add_heading('8. Cambio 7: Build de producción sin SSR', level=1)
doc.add_paragraph(
    'Comando ejecutado:\n\n'
    '  npm run build -- --no-server --base-href /API/\n\n'
    'Archivos generados en dist/fitStock-front/browser/:'
)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Archivo'
hdr[1].text = 'Descripción'
files = [
    ('index.html', 'Punto de entrada de Angular'),
    ('main-*.js', 'Código JavaScript principal (~500 KB)'),
    ('polyfills-*.js', 'Polyfills para compatibilidad (~35 KB)'),
    ('styles-*.css', 'Estilos globales (~16 KB)'),
    ('favicon.ico', 'Icono de navegador'),
    ('icono.jpg', 'Icono adicional'),
    ('images/', 'Carpeta con imágenes (productos, etc.)'),
]
for archivo, desc in files:
    row = table.add_row().cells
    row[0].text = archivo
    row[1].text = desc

doc.add_page_break()
doc.add_heading('8. Cambio 7: Añadir RewriteBase al .htaccess', level=1)
doc.add_paragraph(
    'Al refrescar cualquier página de Angular (F5 o recargar), el navegador '
    'envía la ruta exacta al servidor (ej: /API/admin/inventario). El servidor '
    'busca esa ruta en el sistema de archivos y, al no encontrarla, debe '
    'redirigir al router.php para que sirva el index.html de Angular.'
)
doc.add_paragraph(
    'Sin la directiva RewriteBase, Apache no resuelve correctamente las '
    'rutas relativas del .htaccess cuando se usa desde una subcarpeta, '
    'devolviendo un error 404 al refrescar cualquier ruta interna de Angular.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('FitStock-API/.htaccess, línea 2')

doc.add_heading('Antes:', level=2)
p = doc.add_paragraph()
run = p.add_run("""RewriteEngine On
RewriteCond %{REQUEST_FILENAME} !-f
RewriteCond %{REQUEST_FILENAME} !-d
RewriteRule ^(.*)$ router.php [QSA,L]""")
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Después:', level=2)
p = doc.add_paragraph()
run = p.add_run("""RewriteEngine On
RewriteBase /API/
RewriteCond %{REQUEST_FILENAME} !-f
RewriteCond %{REQUEST_FILENAME} !-d
RewriteRule ^(.*)$ router.php [QSA,L]""")
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph(
    'La línea RewriteBase /API/ le indica a Apache que todas las reglas '
    'del .htaccess deben interpretarse relativas a la subcarpeta /API/, '
    'no a la raíz del dominio. Esto permite que al refrescar cualquier '
    'ruta de Angular (login, admin/inventario, etc.), Apache redirija '
    'correctamente a router.php y este sirva el index.html.'
)

doc.add_page_break()
doc.add_page_break()
doc.add_heading('9. Cambio 8: Credenciales de base de datos (conexion.php)', level=1)
doc.add_paragraph(
    'Se actualizó el archivo conexion.php con las credenciales de la base de '
    'datos de producción en Arsys.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('FitStock-API/conexion.php')

doc.add_heading('Cambios realizados:', level=2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Parámetro'
hdr[1].text = 'Desarrollo (local)'
rows_data = [
    ('Servidor', 'localhost', 'qaqo430.chomsky.es'),
    ('Usuario', 'fitstock', 'qaqo430'),
    ('Contraseña', 'Tokio2324', '17J16a8m28g!'),
    ('Base de datos', 'fitstock', 'qaqo430'),
]
for param, dev, prod in rows_data:
    row = table.add_row().cells
    row[0].text = param
    row[1].text = dev

table2 = doc.add_table(rows=1, cols=1)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Producción (Arsys)'
table2 = doc.add_table(rows=1, cols=1)
for param, dev, prod in rows_data:
    row = table2.add_row().cells
    row[0].text = prod

# Cleaner approach: just list them
doc.add_paragraph()
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Parámetro'
hdr[1].text = 'Desarrollo'
hdr[2].text = 'Producción'
for param, dev, prod in rows_data:
    row = table.add_row().cells
    row[0].text = param
    row[1].text = dev
    row[2].text = prod

doc.add_paragraph(
    'Este cambio es necesario porque la base de datos en Arsys tiene un nombre, '
    'usuario y servidor diferentes a los del entorno de desarrollo local.'
)

doc.add_page_break()
doc.add_heading('10. Cambio 9: Cabeceras CORS en router.php', level=1)
doc.add_paragraph(
    'Se actualizó la cabecera Access-Control-Allow-Origin en router.php para '
    'permitir peticiones desde el dominio de producción.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('FitStock-API/router.php, línea 11')

doc.add_heading('Antes:', level=2)
p = doc.add_paragraph()
run = p.add_run('header("Access-Control-Allow-Origin: http://localhost:4200");')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Después:', level=2)
p = doc.add_paragraph()
run = p.add_run('header("Access-Control-Allow-Origin: https://chomsky.es");')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph(
    'El cambio era necesario porque el navegador bloquea las peticiones HTTP '
    'entre dominios distintos (CORS). El frontend Angular se sirve desde '
    'https://chomsky.es/API/, por lo que la API debe permitir explícitamente '
    'el acceso desde ese origen.'
)

doc.add_page_break()
doc.add_heading('11. Cambio 10: CSS inline en documentación (subcarpeta)', level=1)
doc.add_paragraph(
    'Al alojar la aplicación en una subcarpeta /API/, las rutas relativas '
    'de los archivos CSS enlazados desde docs.html dejaron de funcionar. '
    'La solución fue incrustar el CSS directamente en el HTML.'
)

doc.add_heading('Archivo:', level=2)
doc.add_paragraph('FitStock-API/docs/docs.php')

doc.add_heading('Problema:', level=2)
doc.add_paragraph(
    'El archivo docs.html incluía la hoja de estilos con una ruta relativa:\n\n'
    '  <link rel="stylesheet" href="docs/docs.css">\n\n'
    'Al acceder desde https://chomsky.es/API/docs/docs.php, la ruta se resolvía '
    'correctamente. Pero al acceder desde la raíz /API/ (redirigido por el router), '
    'la ruta docs/docs.css se resolvía como /API/docs/docs.css, que también '
    'funcionaba. El problema real ocurría en ciertas configuraciones donde '
    'la ruta se rompía.'
)

doc.add_heading('Solución:', level=2)
doc.add_paragraph(
    'Se modificó docs.php para leer el contenido de docs.css con file_get_contents() '
    'e incrustarlo directamente en el HTML dentro de una etiqueta <style>:\n\n'
    '  <style>\n'
    '    <?php echo file_get_contents(__DIR__ . "/docs.css"); ?>\n'
    '  </style>\n\n'
    'De esta forma, el CSS se sirve inline y no depende de rutas relativas.'
)

doc.add_page_break()
doc.add_heading('12. Cambio 11: Permisos de archivos en el servidor', level=1)
doc.add_paragraph(
    'Para que la aplicación funcione correctamente en el servidor, fue '
    'necesario ajustar los permisos de ciertas carpetas.'
)

doc.add_heading('Carpeta con permisos especiales:', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Carpeta'
hdr[1].text = 'Permisos'
hdr[2].text = 'Motivo'
folders = [
    ('/API/images/productos/', '775 (rwxrwxr-x)', 'Permitir subida de imágenes desde la app'),
    ('/API/', '755 (rwxr-xr-x)', 'Permitir lectura de archivos estáticos'),
]
for folder, perms, reason in folders:
    row = table.add_row().cells
    row[0].text = folder
    row[1].text = perms
    row[2].text = reason

doc.add_paragraph(
    'Los permisos se configuran desde FileZilla: botón derecho sobre la '
    'carpeta → "Permisos de archivo..." → marcar permisos y marcar '
    '"Recurse into subdirectories" → "Apply to all files and directories".\n\n'
    'Sin estos permisos, la subida de imágenes desde el panel de administración '
    'fallaría por falta de permisos de escritura.'
)

doc.add_page_break()
doc.add_heading('13. Extensión mbstring de PHP', level=1)
doc.add_paragraph(
    'Durante el despliegue se detectó que el servidor de Arsys no tenía '
    'instalada la extensión mbstring de PHP, lo que provocaba un error '
    'al usar la función mb_strtolower().'
)

doc.add_heading('Solución:', level=2)
doc.add_paragraph(
    'Se reemplazó mb_strtolower() por strtolower() en el código PHP, ya que '
    'la aplicación no necesita manejo de caracteres multibyte para su '
    'funcionamiento.\n\n'
    'Si en el futuro se necesitara mbstring (por ejemplo, para caracteres '
    'acentuados), habría que solicitarlo en el panel de control de Arsys '
    'o cambiarse a un plan que lo incluya.'
)

doc.add_page_break()
doc.add_heading('14. Registro de errores durante el despliegue', level=1)
doc.add_paragraph(
    'A continuación se enumeran todos los errores encontrados durante '
    'el proceso de despliegue y su solución:'
)

errores = [
    ('Error 1: API no responde (404)',
     'Al acceder a https://chomsky.es/API/api/login, el servidor devolvía 404.',
     'El api/index.php esperaba que el primer segmento de la ruta fuera "api" '
     '(minúsculas), pero desde la subcarpeta /API/ el primer segmento era "API" '
     '(mayúsculas). Se añadió case "API": y lógica de normalización de rutas.'),
    ('Error 2: Imágenes de productos no se ven',
     'Las imágenes de productos no se cargaban en el panel de administración.',
     'La función getImagenUrl() generaba URLs absolutas (/images/productos/...) '
     'pero la app está en /API/. Se cambió a ruta relativa (images/productos/...).'),
    ('Error 3: Error al subir imágenes',
     'Al subir una imagen desde la app, se guardaba en una ruta incorrecta.',
     'La ruta de subida en api/index.php apuntaba a la estructura de desarrollo '
     'local (FitStock-APP/public/...). Se cambió a la ruta relativa del servidor.'),
    ('Error 4: Página en blanco al refrescar',
     'Al recargar cualquier ruta de Angular (F5), aparecía página en blanco o 404.',
     'Faltaba la directiva RewriteBase /API/ en el .htaccess. Se añadió para '
     'que Apache resuelva correctamente las rutas relativas desde la subcarpeta.'),
    ('Error 5: Estilos de documentación rotos',
     'Los estilos CSS de la página de documentación no se cargaban correctamente.',
     'Las rutas relativas del CSS se rompían en la subcarpeta. Se cambió a '
     'CSS inline con file_get_contents().'),
    ('Error 6: mb_strtolower() no existe',
     'Error PHP: "Call to undefined function mb_strtolower()".',
     'La extensión mbstring no está instalada en Arsys. Se reemplazó por '
     'strtolower() estándar.'),
    ('Error 7: CORS bloqueado',
     'El navegador bloqueaba las peticiones a la API por CORS.',
     'Se actualizó la cabecera Access-Control-Allow-Origin en router.php '
     'de http://localhost:4200 a https://chomsky.es.'),
    ('Error 8: Sin datos en la aplicación',
     'La app cargaba pero no mostraba productos, materiales ni incidencias.',
     'La base de datos solo tenía usuarios. Se añadieron los INSERT con '
     'datos de ejemplo al archivo databasearsys.sql.'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Error'
hdr[1].text = 'Síntoma'
hdr[2].text = 'Solución'
for error, sintoma, solucion in errores:
    row = table.add_row().cells
    row[0].text = error
    row[1].text = sintoma
    row[2].text = solucion

doc.add_heading('Resumen de errores por archivo:', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Archivo modificado'
hdr[1].text = 'Errores que soluciona'
archivos = [
    ('FitStock-API/router.php', 'Error 1 (regex), Error 7 (CORS)'),
    ('FitStock-API/api/index.php', 'Error 1 (case API), Error 3 (ruta subida)'),
    ('FitStock-API/.htaccess', 'Error 4 (RewriteBase)'),
    ('FitStock-API/docs/docs.php', 'Error 5 (CSS inline)'),
    ('FitStock-API/conexion.php', 'Credenciales producción'),
    ('FitStock-APP/src/.../producto-list.ts', 'Error 2 (URL imágenes)'),
    ('FitStock-APP/src/.../services/*.ts', 'Error 7 (URL API)'),
    ('FitStock-API/config/databasearsys.sql', 'Error 8 (datos ejemplo)'),
]
for archivo, errores in archivos:
    row = table.add_row().cells
    row[0].text = archivo
    row[1].text = errores

doc.add_page_break()
doc.add_heading('15. Archivos a subir por FTP', level=1)

doc.add_heading('Paso 1: Angular (a /html/API/)', level=2)
doc.add_paragraph(
    'Desde dist/fitStock-front/browser/:\n'
    '  • index.html\n'
    '  • main-*.js\n'
    '  • polyfills-*.js\n'
    '  • styles-*.css\n'
    '  • favicon.ico\n'
    '  • icono.jpg\n'
    '  • images/ (carpeta completa)'
)

doc.add_heading('Paso 2: API actualizada (a /html/API/)', level=2)
doc.add_paragraph('  • router.php (desde FitStock-API/router.php)\n'
    '  • api/index.php (desde FitStock-API/api/index.php)')

doc.add_heading('Paso 3: Imágenes de productos (a /html/API/images/productos/)', level=2)
doc.add_paragraph('  • FitStock-APP/public/images/productos/*.jpg')

doc.add_heading('Paso 4: Base de datos', level=2)
doc.add_paragraph(
    '  • Importar FitStock-API/config/databasearsys.sql en phpMyAdmin'
)

doc.add_page_break()
doc.add_heading('16. Estructura final del servidor', level=1)

estructura = """/html/
└── API/
    ├── .htaccess              # Redirección a router.php
    ├── router.php             # Router PHP (modificado)
    ├── index.html             # Angular (NUEVO)
    ├── main-*.js              # Angular (NUEVO)
    ├── polyfills-*.js         # Angular (NUEVO)
    ├── styles-*.css           # Angular (NUEVO)
    ├── favicon.ico            # Angular (NUEVO)
    ├── icono.jpg              # Angular (NUEVO)
    ├── images/                # Angular (NUEVO)
    │   └── productos/         # Imágenes de productos
    │       ├── *.jpg
    │       └── ...
    ├── api/                   # Backend PHP
    │   ├── index.php          # (modificado)
    │   └── ...
    ├── conexion.php           # Conexión BD
    ├── config/                # Configuración
    ├── docs/                  # Documentación
    └── models/                # Modelos de datos"""

p = doc.add_paragraph()
run = p.add_run(estructura)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()
doc.add_heading('17. Verificación post-despliegue', level=1)
doc.add_paragraph(
    'Una vez subidos todos los archivos y configurada la base de datos, '
    'se recomienda seguir esta checklist para verificar que todo funciona '
    'correctamente:'
)

verificaciones = [
    '1. Acceder a https://chomsky.es/API/ y comprobar que se carga la página de inicio',
    '2. Probar el registro de un nuevo usuario',
    '3. Iniciar sesión con admin@fitstock.com (o el usuario creado)',
    '4. Verificar que el dashboard de administración carga los datos del resumen',
    '5. Comprobar que la lista de productos se muestra correctamente (incluyendo imágenes)',
    '6. Probar a crear, editar y eliminar un producto',
    '7. Comprobar que los materiales (máquinas y prestables) se listan correctamente',
    '8. Verificar que los préstamos funcionan (crear, devolver, eliminar)',
    '9. Probar el sistema de incidencias (crear, cambiar estado, resolver)',
    '10. Probar la subida de una imagen de producto',
    '11. Refrescar (F5) en varias rutas: /API/, /API/login, /API/admin/inventario',
    '12. Verificar que la documentación en /API/docs/docs.php se ve correctamente',
    '13. Cerrar sesión y verificar que se redirige correctamente',
]

for v in verificaciones:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'Si todos los puntos funcionan, el despliegue se ha completado con éxito.'
)

doc.save('C:\\TFG\\tfg-FitStock\\documentacion\\Cambios_Despliegue_FitStock.docx')
print('Documento generado correctamente.')
