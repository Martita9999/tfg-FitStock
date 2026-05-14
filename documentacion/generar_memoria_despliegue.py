from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ========== PORTADA ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MEMORIA DE CAMBIOS\nDespliegue FitStock')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Migración de Angular SSR a SPA Estático\nConfiguración de API en Producción')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_page_break()

# ========== ÍNDICE ==========
doc.add_heading('Índice', level=1)
items = [
    '1. Resumen de los cambios realizados',
    '2. Cambio 1: URLs de los servicios Angular',
    '3. Cambio 2: Eliminación de SSR (Server-Side Rendering)',
    '4. Cambio 3: Actualización del router.php',
    '5. Cambio 4: Build de producción sin SSR',
    '6. Archivos a subir por FTP',
    '7. Estructura final del servidor',
]
for item in items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ========== 1. RESUMEN ==========
doc.add_heading('1. Resumen de los cambios realizados', level=1)
doc.add_paragraph(
    'Se han realizado los siguientes cambios en el proyecto FitStock para '
    'desplegarlo correctamente en el hosting de Arsys (plan compartido Linux):'
)

cambios = [
    ('Actualización de URLs de API',
     'Todos los servicios Angular apuntaban a http://localhost:8000/api '
     '(entorno de desarrollo). Se ha cambiado a https://chomsky.es/api '
     '(entorno de producción).'),
    ('Desactivación de SSR',
     'La aplicación usaba Angular SSR (Server-Side Rendering), que requiere '
     'Node.js en el servidor. Arsys (plan compartido) no soporta Node.js, '
     'por lo que se ha compilado la aplicación como SPA estático.'),
    ('Actualización del router.php',
     'El router de la API se ha modificado para que sirva el index.html '
     'de Angular en todas las rutas que no sean de la API, permitiendo '
     'el enrutamiento interno de Angular (SPA).'),
    ('Build de producción',
     'Se ha ejecutado ng build --no-server para generar los archivos '
     'estáticos optimizados para producción.'),
]
for titulo, desc in cambios:
    p = doc.add_paragraph()
    run = p.add_run(f'{titulo}: ')
    run.bold = True
    p.add_run(desc)

# ========== 2. URLs ==========
doc.add_page_break()
doc.add_heading('2. Cambio 1: URLs de los servicios Angular', level=1)
doc.add_paragraph(
    'Todos los servicios de Angular que se comunican con la API tenían '
    'la URL http://localhost:8000/api (entorno de desarrollo local con '
    'PHP Artisan Serve). Se ha cambiado a https://chomsky.es/api para '
    'que apunten al servidor de producción.'
)

doc.add_heading('Archivos modificados (7 servicios):', level=2)

servicios = [
    ('src/app/services/compras.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/incidencias.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/materiales.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/prestamos.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/productos.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/resumen.service.ts', 'private API_URL = \'http://localhost:8000/api\''),
    ('src/app/services/usuario.ts', 'private API_URL = \'http://localhost:8000/api\''),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Archivo'
hdr[1].text = 'Línea modificada'
for archivo, linea in servicios:
    row = table.add_row().cells
    row[0].text = archivo
    row[1].text = linea

doc.add_paragraph()
doc.add_paragraph(
    'Cambio realizado: se sustituyó http://localhost:8000/api por '
    'https://chomsky.es/api en cada uno de estos archivos.'
)

# ========== 3. SSR ==========
doc.add_page_break()
doc.add_heading('3. Cambio 2: Eliminación de SSR', level=1)
doc.add_paragraph(
    'La aplicación Angular estaba configurada con SSR (Server-Side Rendering), '
    'que genera una versión de la app en Node.js para renderizar en servidor. '
    'Como el hosting compartido de Arsys no ejecuta Node.js, se ha optado por '
    'compilar la aplicación como SPA estático (solo archivos HTML, CSS y JS).'
)

doc.add_heading('¿Qué es SSR y por qué no lo necesitas?', level=2)
doc.add_paragraph(
    'SSR sirve para que los motores de búsqueda (Google) puedan indexar '
    'el contenido de la página y para mejorar la velocidad percibida en la '
    'primera carga. Como FitStock requiere inicio de sesión (login), '
    'el SEO es irrelevante. La app funcionará exactamente igual sin SSR.'
)

doc.add_heading('¿Cómo se ha desactivado?', level=2)
doc.add_paragraph(
    'Se ha ejecutado el build con la flag --no-server:\n\n'
    '  ng build --no-server\n\n'
    'Esto produce solo los archivos del navegador (browser/) sin generar '
    'la parte del servidor (server/). No ha sido necesario modificar '
    'código fuente ni eliminar archivos de configuración de SSR.'
)

# ========== 4. router.php ==========
doc.add_page_break()
doc.add_heading('4. Cambio 3: Actualización del router.php', level=1)
doc.add_paragraph(
    'El router.php es el punto de entrada de todas las peticiones HTTP '
    'al servidor. Originalmente solo manejaba la API y la documentación. '
    'Se ha modificado para que también sirva el index.html de Angular '
    'en todas las rutas que no pertenezcan a la API.'
)

doc.add_heading('Código original:', level=2)
original = """<?php
header("Access-Control-Allow-Origin: https://chomsky.es");
// ... cabeceras CORS ...

$uri = parse_url($_SERVER['REQUEST_URI'], PHP_URL_PATH);

// Servir archivos estáticos
$filePath = __DIR__ . $uri;
if ($uri !== '/' && is_file($filePath)) {
    return false;
}

// Redirigir /api al backend
if (preg_match('#^/api#', $uri)) {
    require __DIR__ . '/api/index.php';
    return true;
}

// Fallback: mostrar documentación
if (is_file(__DIR__ . '/docs/docs.php')) {
    require __DIR__ . '/docs/docs.php';
} else {
    echo "API FitStock activa.";
}
return true;"""

p = doc.add_paragraph()
run = p.add_run(original)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Código modificado:', level=2)
modificado = """<?php
header("Access-Control-Allow-Origin: https://chomsky.es");
// ... cabeceras CORS ...

$uri = parse_url($_SERVER['REQUEST_URI'], PHP_URL_PATH);

// Servir archivos estáticos (CSS, JS, imágenes...)
$filePath = __DIR__ . $uri;
if ($uri !== '/' && is_file($filePath)) {
    return false;
}

// Redirigir /api al backend
if (preg_match('#^/api#', $uri)) {
    require __DIR__ . '/api/index.php';
    return true;
}

// NUEVO: Servir index.html de Angular para SPA
$angularIndex = __DIR__ . '/index.html';
if (is_file($angularIndex)) {
    http_response_code(200);
    header("Content-Type: text/html; charset=utf-8");
    require $angularIndex;
    return true;
}

// Fallback: documentación
if (is_file(__DIR__ . '/docs/docs.php')) {
    require __DIR__ . '/docs/docs.php';
} else {
    echo "API FitStock activa.";
}
return true;"""

p = doc.add_paragraph()
run = p.add_run(modificado)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Explicación del cambio:', level=2)
doc.add_paragraph(
    'La clave está en la sección "NUEVO": después de comprobar que la petición '
    'no es a la API ni a un archivo estático existente, el router busca '
    'el archivo index.html de Angular en la raíz del servidor. Si existe, '
    'lo sirve. Esto permite que rutas como /login, /admin/inventario, etc. '
    'funcionen correctamente gracias al enrutamiento interno de Angular (SPA). '
    'Si no existe index.html, muestra la documentación como fallback.'
)

# ========== 5. BUILD ==========
doc.add_page_break()
doc.add_heading('5. Cambio 4: Build de producción sin SSR', level=1)
doc.add_paragraph(
    'Se ha ejecutado el siguiente comando para compilar la aplicación Angular '
    'en modo producción, generando únicamente los archivos estáticos del '
    'navegador (sin la parte del servidor Node.js):'
)

p = doc.add_paragraph()
run = p.add_run('\n  npm run build -- --no-server\n')
run.font.name = 'Consolas'
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'Este comando genera los archivos de salida en la carpeta:\n'
    '  FitStock-APP/dist/fitStock-front/browser/\n\n'
    'Contenido generado:'
)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Archivo'
hdr[1].text = 'Descripción'
files = [
    ('index.html', 'Punto de entrada de la aplicación Angular'),
    ('main-H27MESBC.js', 'Código JavaScript principal (500 KB)'),
    ('polyfills-DOYHMSTV.js', 'Polyfills para compatibilidad (35 KB)'),
    ('styles-RXQRMPP5.css', 'Estilos globales de la aplicación (16 KB)'),
    ('favicon.ico', 'Icono de la pestaña del navegador'),
    ('icono.jpg', 'Icono adicional de la aplicación'),
    ('images/', 'Carpeta con imágenes utilizadas en la app'),
]
for archivo, desc in files:
    row = table.add_row().cells
    row[0].text = archivo
    row[1].text = desc

# ========== 6. FTP ==========
doc.add_page_break()
doc.add_heading('6. Archivos a subir por FTP', level=1)
doc.add_paragraph(
    'Para completar el despliegue, se deben subir los siguientes archivos '
    'al servidor a través de Filezilla (o cualquier cliente FTP):'
)

doc.add_heading('Paso 1: Subir los archivos de Angular', level=2)
doc.add_paragraph(
    'Desde la carpeta local:\n'
    '  FitStock-APP/dist/fitStock-front/browser/\n\n'
    'A la raíz del servidor (donde ya están api/, router.php, etc.):\n\n'
    'Archivos a subir:\n'
    '  • index.html\n'
    '  • main-H27MESBC.js\n'
    '  • polyfills-DOYHMSTV.js\n'
    '  • styles-RXQRMPP5.css\n'
    '  • favicon.ico\n'
    '  • icono.jpg\n'
    '  • images/ (carpeta completa)'
)

doc.add_heading('Paso 2: Subir el router.php actualizado', level=2)
doc.add_paragraph(
    'Desde la carpeta local:\n'
    '  FitStock-API/router.php\n\n'
    'A la raíz del servidor, SOBRESCRIBIENDO el existente.'
)

doc.add_heading('Paso 3: Opcional - Subir documentación', level=2)
doc.add_paragraph(
    'Si se desea mantener la documentación, asegurarse de que la carpeta '
    'docs/ está presente en el servidor.'
)

# ========== 7. ESTRUCTURA ==========
doc.add_page_break()
doc.add_heading('7. Estructura final del servidor', level=1)
doc.add_paragraph(
    'Tras el despliegue, la estructura en la raíz del servidor debería ser '
    'la siguiente:'
)

estructura = """/
├── .htaccess            # Redirección a router.php
├── router.php           # Router principal (PHP) - ACTUALIZADO
├── index.html           # Angular - NUEVO
├── main-H27MESBC.js     # Angular - NUEVO
├── polyfills-DOYHMSTV.js # Angular - NUEVO
├── styles-RXQRMPP5.css  # Angular - NUEVO
├── favicon.ico          # Angular - NUEVO
├── icono.jpg            # Angular - NUEVO
├── images/              # Angular - NUEVO
├── api/                 # Backend PHP
│   ├── index.php
│   └── ...
├── conexion.php         # Conexión a base de datos
├── config/              # Configuración
├── docs/                # Documentación
└── models/              # Modelos de datos"""

p = doc.add_paragraph()
run = p.add_run(estructura)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ========== GUARDAR ==========
doc.save('C:\\TFG\\tfg-FitStock\\documentacion\\Cambios_Despliegue_FitStock.docx')
print('Documento generado correctamente.')
