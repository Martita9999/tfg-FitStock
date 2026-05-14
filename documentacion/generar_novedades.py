from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# ── Estilos ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    h.font.bold = True

# ── Portada ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FitStock — Documentación de Novedades')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Detalles técnicos de los nuevos cambios implementados')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mayo 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# ── Función para añadir con formato de código ──
def add_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ── Contenido ──
doc.add_heading('1. Documentación de la API (carpeta /docs/)', level=1)

doc.add_heading('1.1 docs/docs.php — Generador dinámico', level=2)

doc.add_paragraph('Archivo PHP que contiene la lógica para generar la página de documentación de la API. Estructura:')

add_bullet('Array $resources (líneas 2-72) que define todos los endpoints agrupados por recurso. Cada grupo contiene:', bold_prefix='$resources: ')
add_bullet('base: ruta base del recurso (ej. /api/usuarios)')
add_bullet('endpoints: array de [Método HTTP, Ruta, Descripción, Clave de autenticación]')
add_bullet('Grupos definidos: Autenticación, Usuarios, Materiales, Préstamos, Incidencias, Productos, Compras, Resumen')

add_bullet('Array $authLabels (líneas 74-79) que mapea claves de autenticación a [etiqueta, clase CSS]: any → Sin auth, all → Autenticado, admin → Admin, admin-ent → Admin/Entrenador', bold_prefix='$authLabels: ')

add_bullet('Array $methodColors (línea 81): GET → azul, POST → verde, PUT → ámbar, DELETE → rojo', bold_prefix='$methodColors: ')

add_bullet('Generación del HTML (líneas 83-102): Itera sobre cada recurso y endpoint construyendo el HTML con clases CSS.', bold_prefix='Bucle de renderizado: ')

add_bullet('Carga docs.html como plantilla (línea 104): $html = file_get_contents(__DIR__ . \'/docs.html\')', bold_prefix='Plantilla: ')

add_bullet('Reemplazo del marcador (línea 105): echo str_replace(\'{{RESOURCES}}\', $rows, $html)', bold_prefix='Inserción: ')

doc.add_heading('1.2 docs/docs.html — Plantilla HTML', level=2)

doc.add_paragraph('Plantilla HTML5 minimalista con:')
add_bullet('<meta charset="UTF-8"> y viewport responsive')
add_bullet('<link rel="stylesheet" href="/docs/docs.css"> — enlace al CSS')
add_bullet('{{RESOURCES}} — marcador de posición que docs.php reemplaza con los endpoints')
add_bullet('Footer con texto "FitStock API — Backend del sistema de gestión"')

doc.add_heading('1.3 docs/docs.css — Estilos', level=2)

doc.add_paragraph('Hoja de estilos con formato expandido (cada propiedad en su línea). Incluye:')
add_bullet('Reset universal: margin:0, padding:0, box-sizing:border-box')
add_bullet('Body con gradiente claro #f8fafc → #f1f5f9, fuente Segoe UI')
add_bullet('.resource: tarjetas blancas con border-radius 14px, sombra sutil, borde 1px #e2e8f0')
add_bullet('.method: badges de método HTTP con clases de color .get (azul), .post (verde), .put (ámbar), .delete (rojo)')
add_bullet('.route: ruta en monospace (Cascadia Code / Fira Code)')
add_bullet('.auth: badges de autenticación .auth-admin (rojo), .auth-all (azul), .auth-any (gris)')
add_bullet('Media query a 768px para responsive: los endpoints se envuelven')

doc.add_page_break()

doc.add_heading('2. Router (router.php)', level=1)

doc.add_paragraph('Cambios realizados en el archivo de enrutamiento:')

add_bullet('Líneas 14-18: Nueva funcionalidad para servir archivos estáticos')
add_code('$filePath = __DIR__ . $uri;')
add_code('if ($uri !== \'/\' && is_file($filePath)) {')
add_code('    return false;')
add_code('}')
add_bullet('Si la URI no es la raíz y el archivo existe en el sistema de archivos, devuelve false para que el servidor PHP built-in sirva el archivo con su MIME type correcto.')
add_bullet('Esto permite que el navegador cargue docs/docs.css correctamente.')
add_bullet('Las rutas /api siguen delegando en api/index.php.')
add_bullet('Cualquier otra ruta muestra la documentación HTML mediante require de docs/docs.php.')

doc.add_heading('3. Frontend — Modo Oscuro', level=1)

doc.add_heading('3.1 index.html — Inicialización temprana', level=2)
add_bullet('Script inline en <head> que lee localStorage antes de que Angular cargue:')
add_code('if (localStorage.getItem(\'darkMode\') === \'true\') {')
add_code('  document.documentElement.setAttribute(\'data-theme\', \'dark\');')
add_code('}')
add_bullet('Soluciona el flash de modo claro al recargar la página con modo oscuro activo.')

doc.add_heading('3.2 admin-sidebar.ts — applyTheme()', level=2)
add_bullet('Modo oscuro: setAttribute(\'data-theme\', \'dark\')')
add_bullet('Modo claro: removeAttribute(\'data-theme\') en lugar de setAttribute(\'data-theme\', \'light\')')
add_bullet('Esto evita tener que definir un bloque CSS [data-theme="light"], ya que :root contiene los valores de modo claro por defecto.')

doc.add_heading('3.3 login.ts — Toggle en login', level=2)
add_bullet('Constructor: inicializa darkMode desde localStorage y aplica el tema.')
add_bullet('toggleDarkMode(): invierte el estado, persiste en localStorage y actualiza data-theme.')

doc.add_heading('3.4 login.html — Botón de tema', level=2)
add_bullet('Añadido <button class="theme-toggle"> en la esquina superior derecha.')
add_bullet('Muestra 🌙 en modo claro y ☀️ en modo oscuro.')

doc.add_heading('3.5 styles.css — Variables de modo claro mejoradas', level=2)
add_bullet('--bg: #f8fafc → #f1f5f9 (más gris para contraste con tarjetas blancas)')
add_bullet('--border: #e2e8f0 → #cbd5e1 (bordes más visibles)')
add_bullet('Nuevas variables: --error-bg, --error-color, --error-border, --success-bg, --success-color, --success-border, --warning-bg, --warning-color')
add_bullet('Cada variable tiene valores distintos para :root (modo claro) y [data-theme="dark"]')

doc.add_heading('3.6 Inputs — Color de texto en modo oscuro', level=2)
add_bullet('Se añadió color: var(--text) y background: var(--card) a los inputs en: login.css, registro.css, usuario-list.css, dashboard-home.css (.input-sm)')
add_bullet('Esto asegura que el texto sea blanco (#f1f5f9) sobre fondo oscuro (#1e293b) en modo oscuro.')

doc.add_page_break()

doc.add_heading('4. Frontend — Dashboard', level=1)

doc.add_heading('4.1 dashboard-home.ts — Nuevos datos', level=2)
add_bullet('prestamosPendientes: array de préstamos sin devolución (filtrados con !p.devolucion)')
add_bullet('totalUsuarios: número total de usuarios registrados')
add_bullet('En cargarResumen(): dos nuevas llamadas a prestamosService.getPrestamos() y usuarioService.getUsuarios()')

doc.add_heading('4.2 dashboard-home.html — Nuevas tarjetas', level=2)
add_bullet('Tarjeta "🤝 Préstamos Pendientes": muestra contador y lista hasta 5 préstamos con usuario y material')
add_bullet('Tarjeta "👥 Usuarios": muestra el total de usuarios registrados')
add_bullet('Ambas tarjetas enlazan a sus respectivas secciones (/admin/prestamos, /admin/usuarios)')

doc.add_heading('5. Varios Frontend', level=1)

doc.add_heading('5.1 usuario-list.html — 🔑 para entrenadores', level=2)
add_bullet('Condición del botón 🔑 cambiada para permitir que admin solicite cambio de contraseña a entrenadores.')
add_bullet('Antes: u.rol === \'cliente\'')
add_bullet('Ahora: (userRole === \'admin\' || u.rol === \'cliente\') && u.rol !== \'admin\'')

doc.add_heading('5.2 login.ts — Forzar cambio para entrenadores', level=2)
add_bullet('Condición del modal de cambio de contraseña obligatorio ampliada:')
add_bullet('Antes: res.user?.rol === \'cliente\'')
add_bullet('Ahora: (res.user?.rol === \'cliente\' || res.user?.rol === \'entrenador\')')

doc.add_heading('5.3 prestamo-list.html — Devuelto para todos', level=2)
add_bullet('Se eliminó el filtro && userRole === \'cliente\' del botón "Devuelto"')
add_bullet('Ahora admin y entrenador también pueden marcar préstamos como devueltos directamente')

doc.add_heading('5.4 producto-list.html — Texto Stock', level=2)
add_bullet('Cambió \'Stock OK\' por \'Stock\' (texto más limpio)')

doc.add_heading('5.5 admin-dashboard.css — Fondo adaptable', level=2)
add_bullet('.dashboard-wrapper: de gradiente fijo a var(--bg-gradient)')
add_bullet('.data-table tbody tr:hover: de #fff7ed a var(--bg)')

doc.add_heading('6. Backend — Usuario.php', level=1)

doc.add_heading('6.1 obtenerSiguienteIdLibre() — Nuevo método', level=2)
add_bullet('Obtiene todos los IDs existentes ordenados')
add_bullet('Itera desde 1 buscando el primer hueco (gap)')
add_bullet('Devuelve el primer ID libre para reutilizar')

doc.add_heading('6.2 crear() — Modificado', level=2)
add_bullet('Antes: INSERT sin id_usuario (AUTO_INCREMENT)')
add_bullet('Ahora: calcula el siguiente ID libre y lo inserta explícitamente')
add_bullet('Permite reutilizar IDs de usuarios eliminados')
add_code('$nuevoId = self::obtenerSiguienteIdLibre();')
add_code('$stmt = $conexion->prepare("INSERT INTO usuarios (id_usuario, nombre, email, password_hash, rol) VALUES (?, ?, ?, ?, ?)");')

# ── Guardar ──
doc.save('C:\\TFG\\tfg-FitStock\\docs-novedades.docx')
print("Word document created successfully")
