from docx import Document

doc = Document()
doc.add_heading("Problema de rutas CSS en produccion (Arsys)", level=1)

doc.add_heading("Problema", level=2)
doc.add_paragraph(
    "Al acceder a router.php directamente desde una subcarpeta (/API/), "
    "la ruta relativa del CSS no se resolvia correctamente porque el navegador "
    "calcula las rutas relativas respecto a la URL de la pagina, no respecto "
    "a la ubicacion real del archivo PHP."
)

doc.add_heading("Solucion", level=2)
p = doc.add_paragraph()
p.add_run("En lugar de un <link> externo, se incrusto el CSS directamente en el HTML mediante PHP ")
r = p.add_run("file_get_contents + str_replace")
r.bold = True
p.add_run(", eliminando cualquier dependencia de rutas.")

doc.add_heading("Detalle tecnico", level=2)
doc.add_paragraph(
    "Originalmente docs.html tenia <link rel='stylesheet' href='/docs/docs.css'>. "
    "Al estar la API dentro de /API/, la ruta absoluta /docs/docs.css no coincidia "
    "con la ubicacion real /API/docs/docs.css. Se intento usar rutas relativas "
    "(docs/docs.css) pero el navegador las resuelve respecto a la URL de la pagina, "
    "no respecto a la carpeta del archivo PHP."
)
doc.add_paragraph(
    "La solucion definitiva fue modificar docs.php para leer el contenido de docs.css "
    "con file_get_contents() e inyectarlo directamente en el HTML mediante un placeholder "
    "{{CSS}} dentro de una etiqueta <style>. Asi el CSS viaja incrustado en la propia "
    "respuesta HTML sin necesidad de peticiones adicionales ni rutas."
)

doc.add_heading("Codigo aplicado", level=2)
doc.add_paragraph(
    "En docs.php:\n"
    "$css = file_get_contents(__DIR__ . '/docs.css');\n"
    "echo str_replace(['{{RESOURCES}}', '{{CSS}}'], [$rows, $css], $html);\n\n"
    "En docs.html:\n"
    "<style>{{CSS}}</style>"
)

doc.save("documentacion/PROBLEMA_RUTAS_CSS.docx")
print("Documento creado correctamente")
