from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Configuración de estilos --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# -- Título principal --
title = doc.add_heading('Análisis de Dependencias del Proyecto FitStock', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Documento explicativo de las dependencias definidas en package.json '
    'del frontend Angular (FitStock-APP) y su importancia en el proyecto.'
)

doc.add_paragraph('')

# -- 1. Introducción --
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'El proyecto FitStock-APP es una aplicación frontend construida con Angular 21. '
    'El archivo package.json define todas las dependencias necesarias para desarrollar, '
    'compilar, probar y ejecutar la aplicación. A continuación se detalla cada dependencia '
    'y su función dentro del proyecto.'
)

# -- 2. Dependencias de producción --
doc.add_heading('2. Dependencias de Producción (dependencies)', level=1)
doc.add_paragraph(
    'Son las librerías necesarias para que la aplicación funcione en tiempo de ejecución, '
    'tanto en el navegador como en el servidor (SSR).'
)

prod_deps = [
    {
        'name': '@angular/core',
        'why': 'Es el núcleo del framework Angular. Proporciona los mecanismos fundamentales: '
               'componentes, directivas, servicios, inyección de dependencias, detección de cambios '
               'y ciclo de vida de la aplicación. Sin ella no existe la aplicación.',
        'importance': 'Crítica'
    },
    {
        'name': '@angular/common',
        'why': 'Contiene las directivas y pipes más utilizadas como ngIf, ngFor, ngClass, ngStyle, '
               'CurrencyPipe, DatePipe, etc. También incluye el servicio HttpClient para realizar '
               'peticiones HTTP a la API.',
        'importance': 'Crítica'
    },
    {
        'name': '@angular/compiler',
        'why': 'Compila las plantillas (templates) de Angular en código JavaScript ejecutable. '
               'Angular usa un compilador JIT (Just-In-Time) o AOT (Ahead-Of-Time); en producción '
               'se usa AOT para mejor rendimiento.',
        'importance': 'Crítica'
    },
    {
        'name': '@angular/forms',
        'why': 'Permite construir y gestionar formularios en la aplicación. Proporciona dos enfoques: '
               'Template-Driven Forms (mediante directivas en la plantilla) y Reactive Forms '
               '(mediante modelos en TypeScript). Esencial para el registro de usuarios, '
               'inicio de sesión, y cualquier entrada de datos.',
        'importance': 'Alta'
    },
    {
        'name': '@angular/router',
        'why': 'Gestiona la navegación entre las diferentes páginas/vistas de la aplicación '
               '(ej. página de inicio, dashboard, perfil de usuario, etc.). Permite definir rutas, '
               'guardianes (guards) para proteger rutas, y lazy loading de módulos.',
        'importance': 'Alta'
    },
    {
        'name': '@angular/platform-browser',
        'why': 'Proporciona los servicios y componentes específicos para ejecutar Angular en un '
               'navegador web. Incluye el renderizador DOM y el soporte para eventos del navegador.',
        'importance': 'Crítica'
    },
    {
        'name': '@angular/platform-server',
        'why': 'Permite ejecutar Angular en el servidor (Server-Side Rendering). Mejora el SEO '
               'y el tiempo de carga inicial al renderizar la primera vista en el servidor.',
        'importance': 'Media'
    },
    {
        'name': '@angular/ssr',
        'why': 'Paquete oficial de Angular para Server-Side Rendering. Configura y gestiona la '
               'comunicación entre el servidor Express y la aplicación Angular para SSR.',
        'importance': 'Media'
    },
    {
        'name': 'express',
        'why': 'Framework web para Node.js. En este proyecto actúa como servidor HTTP para servir '
               'la aplicación Angular renderizada en el servidor (SSR) y manejar rutas del lado servidor.',
        'importance': 'Media'
    },
    {
        'name': 'rxjs',
        'why': 'Librería de programación reactiva basada en Observables. Angular la utiliza internamente '
               'para manejar eventos asíncronos: peticiones HTTP (HttpClient), cambios en formularios '
               '(valueChanges), eventos de usuario, y comunicación entre componentes mediante servicios.',
        'importance': 'Alta'
    },
    {
        'name': 'tslib',
        'why': 'Proporciona funciones auxiliares de TypeScript que se usan en el código compilado. '
               'Ayuda a reducir el tamaño del bundle al evitar duplicar código de utility.',
        'importance': 'Baja'
    },
    {
        'name': 'zone.js',
        'why': 'Librería que intercepta operaciones asíncronas (setTimeout, promesas, eventos DOM, etc.) '
               'para notificar a Angular cuándo debe ejecutar la detección de cambios y actualizar la UI. '
               'Es fundamental para que Angular "sepa" cuándo ha cambiado el estado.',
        'importance': 'Crítica'
    },
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Dependencia'
hdr[1].text = 'Importancia'
hdr[2].text = 'Explicación'

for dep in prod_deps:
    row = table.add_row().cells
    row[0].text = dep['name']
    row[1].text = dep['importance']
    row[2].text = dep['why']

doc.add_paragraph('')

# -- 3. Dependencias de desarrollo --
doc.add_heading('3. Dependencias de Desarrollo (devDependencies)', level=1)
doc.add_paragraph(
    'Son herramientas necesarias durante el desarrollo, compilación y pruebas, pero no se incluyen '
    'en el bundle de producción.'
)

dev_deps = [
    {
        'name': '@angular/cli',
        'why': 'La Interfaz de Línea de Comandos de Angular. Permite ejecutar comandos como '
               '"ng serve" (servidor de desarrollo), "ng build" (compilar), "ng generate" (generar '
               'componentes, servicios, etc.), "ng test" (ejecutar tests). Es la herramienta principal '
               'del desarrollador.',
        'importance': 'Crítica'
    },
    {
        'name': '@angular/build',
        'why': 'Sistema de compilación optimizado para Angular. Gestiona la transformación de '
               'TypeScript a JavaScript, el empaquetado (bundling), la minificación, y la generación '
               'de assets. Reemplaza a los antiguos build-angular/webpack.',
        'importance': 'Alta'
    },
    {
        'name': '@angular/compiler-cli',
        'why': 'Compilador de Angular para la compilación AOT (Ahead-of-Time). Analiza el código '
               'TypeScript y las plantillas HTML en tiempo de compilación para generar código '
               'JavaScript optimizado y detectar errores antes de la ejecución.',
        'importance': 'Alta'
    },
    {
        'name': 'typescript',
        'why': 'Lenguaje de programación en el que está escrita toda la aplicación. Proporciona '
               'tipado estático, interfaces, decoradores y otras características que mejoran la '
               'calidad del código y la productividad del desarrollador.',
        'importance': 'Crítica'
    },
    {
        'name': 'vitest',
        'why': 'Framework de testing moderno y rápido. Se utiliza para ejecutar pruebas unitarias '
               'de los componentes, servicios y otras piezas de la aplicación. Vitest ofrece una '
               'experiencia similar a Jest pero con mejor integración con Vite/Angular.',
        'importance': 'Alta'
    },
    {
        'name': '@types/jasmine',
        'why': 'Tipados de TypeScript para Jasmine (framework de testing). Proporciona autocompletado '
               'y verificación de tipos al escribir tests con sintaxis Jasmine.',
        'importance': 'Baja'
    },
    {
        'name': '@types/jest',
        'why': 'Tipados de TypeScript para Jest. Aunque el proyecto usa Vitest, muchos tipos y '
               'utility types son compatibles con Jest, lo que facilita la migración o el uso de '
               'librerías que extienden Jest.',
        'importance': 'Baja'
    },
    {
        'name': '@types/express',
        'why': 'Tipados de TypeScript para Express. Permite usar Express con seguridad de tipos, '
               'autocompletado y verificación en tiempo de compilación.',
        'importance': 'Baja'
    },
    {
        'name': '@types/node',
        'why': 'Tipados de TypeScript para Node.js. Necesario para usar APIs de Node.js (fs, path, '
               'process, etc.) con seguridad de tipos, especialmente en el contexto de SSR.',
        'importance': 'Media'
    },
    {
        'name': 'jsdom',
        'why': 'Simula un entorno de navegador en Node.js. Utilizado por Vitest para ejecutar tests '
               'de componentes Angular sin necesidad de un navegador real.',
        'importance': 'Media'
    },
    {
        'name': 'prettier',
        'why': 'Formateador de código automático. Mantiene un estilo de código consistente en todo '
               'el proyecto (indentación, comillas, punto y coma, etc.), evitando discusiones de estilo '
               'y mejorando la legibilidad.',
        'importance': 'Baja'
    },
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Dependencia'
hdr2[1].text = 'Importancia'
hdr2[2].text = 'Explicación'

for dep in dev_deps:
    row = table2.add_row().cells
    row[0].text = dep['name']
    row[1].text = dep['importance']
    row[2].text = dep['why']

doc.add_paragraph('')

# -- 4. packageManager --
doc.add_heading('4. packageManager', level=1)
doc.add_paragraph(
    'El campo "packageManager": "npm@11.11.0" en package.json indica que el proyecto está '
    'configurado para usar npm versión 11.11.0. Esto asegura que todos los desarrolladores usen '
    'la misma versión, evitando inconsistencias (Corepack).'
)

# -- 5. Resumen --
doc.add_heading('5. Resumen de la Importancia', level=1)
doc.add_paragraph(
    'En orden de importancia para que el proyecto funcione correctamente:\n\n'
    '1. @angular/core, @angular/common, @angular/compiler, @angular/platform-browser — '
    'El núcleo de Angular, sin estas no hay aplicación.\n\n'
    '2. zone.js — Necesario para la detección de cambios de Angular.\n\n'
    '3. typescript — Lenguaje base del proyecto.\n\n'
    '4. @angular/cli — Herramienta indispensable para desarrollar y compilar.\n\n'
    '5. @angular/forms y @angular/router — Funcionalidades clave para la interacción del usuario.\n\n'
    '6. rxjs — Programación reactiva usada extensivamente.\n\n'
    '7. vitest y @angular/build — Testing y compilación.\n\n'
    '8. @angular/platform-server, @angular/ssr, express — SSR para SEO y rendimiento.\n\n'
    '9. @types/node, jsdom — Soporte para testing y entorno Node.\n\n'
    '10. tslib, prettier, @types/* — Utilidades que mejoran la experiencia de desarrollo.'
)

# -- Guardar --
output_path = 'C:\\TFG\\tfg-FitStock\\Analisis_Dependencias_FitStock.docx'
doc.save(output_path)
print(f'Documento generado en: {output_path}')
