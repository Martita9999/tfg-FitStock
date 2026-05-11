from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('Reporte de Cambios y Solución de Errores — FitStock', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'Fecha: {date.today().strftime("%d/%m/%Y")}')
doc.add_paragraph('')

# ============================================================
# 1. RESUMEN GENERAL
# ============================================================
doc.add_heading('1. Resumen General', level=1)
doc.add_paragraph(
    'Este documento recoge todos los cambios realizados en el proyecto FitStock, '
    'incluyendo la implementación de nuevas funcionalidades, corrección de errores, '
    'reorganización de archivos y sincronización de la base de datos.'
)

doc.add_paragraph('')

# ============================================================
# 2. CAMBIOS REALIZADOS
# ============================================================
doc.add_heading('2. Cambios Realizados', level=1)

# 2.1 Productos
doc.add_heading('2.1 Página de Productos (Inventario)', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Cambio'
hdr[1].text = 'Archivos afectados'
hdr[2].text = 'Descripción'

rows_data = [
    ['Rediseño de tabla a tarjetas', 'producto-list.html, producto-list.css',
     'Se reemplazó la tabla horizontal por un grid responsivo de tarjetas verticales con imagen, nombre, descripción, precio y stock.'],
    ['Campo descripción en productos', 'API: Producto.php, index.php\nAPP: app.interfaces.ts, productos.service.ts, producto-list.ts',
     'Se añadió la columna "descripcion" a productos_stock y se propagó por toda la pila (PHP → API → Angular).'],
    ['Imágenes de productos', 'public/images/productos/ (8 archivos)\nproducto-list.ts (getImagenUrl)',
     'Las imágenes se movieron de la raíz a public/images/productos/ y se renombraron para coincidir con los nombres en BD. Se usa encodeURIComponent() para manejar espacios.'],
    ['Precio como número', 'API: index.php\nTemplate: producto-list.html',
     'Se devuelve floatval($p->getPrecio()) desde PHP y se usa (+p.precio).toFixed(2) en la plantilla para evitar errores de tipo string.'],
    ['Mensaje de compra exitosa', 'producto-list.ts, producto-list.css',
     'Mensaje verde con gradiente, animación slide-down, centrado, duración de 1 minuto (setTimeout 60000ms).'],
]

for r in rows_data:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph('')

# 2.2 Préstamos
doc.add_heading('2.2 Gestión de Préstamos', level=2)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Cambio'
hdr2[1].text = 'Archivos afectados'
hdr2[2].text = 'Descripción'

rows_data2 = [
    ['Creación de material prestable', 'prestamo-list.ts, prestamo-list.html, prestamo-list.css\nmateriales.service.ts',
     'Botón "Añadir Material" sobre la sección de materiales disponibles. Modal con nombre, descripción y cantidad a crear. Usa MaterialesService.createMaterial() con tipo: "prestable".'],
    ['Botones Editar/Borrar en tarjetas', 'prestamo-list.ts, prestamo-list.html, prestamo-list.css',
     'Botones ✏️ Editar y 🗑️ Borrar debajo del selector de cantidad en cada grupo de material. Editar abre modal para cambiar nombre/descripción (aplica a todas las unidades). Borrar elimina una unidad disponible del grupo.'],
    ['Campo ubicación en materiales', 'database.sql\nMaterial.php',
     'Se añadió columna "ubicacion" a la tabla material.'],
    ['Estados de material ampliados', 'database.sql\nDocker',
     'ENUM estado ampliado a 7 valores: operativo, prestado, en_reparacion, deteriorado, extraviado, de_baja, en_mantenimiento.'],
]

for r in rows_data2:
    row = table2.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph('')

# 2.3 Base de datos
doc.add_heading('2.3 Base de Datos', level=2)

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Cambio'
hdr3[1].text = 'Descripción'

rows_data3 = [
    ['Sincronización database.sql con Docker',
     'Se actualizó database.sql para reflejar exactamente el esquema de Docker: misma estructura de tablas, mismos ENUMs, mismos datos de ejemplo. Se eliminaron ALTER TABLE redundantes que ya no aplicaban.'],
    ['Nuevas columnas',
     'descripcion en productos_stock, ubicacion en material.'],
    ['Nuevos valores ENUM',
     'material.estado ampliado de 3 a 7 valores.'],
]

for r in rows_data3:
    row = table3.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

doc.add_paragraph('')

# 2.4 Limpieza
doc.add_heading('2.4 Limpieza de Archivos', level=2)

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Archivo'
hdr4[1].text = 'Acción'
hdr4[2].text = 'Motivo'

rows_data4 = [
    ['10 JPG duplicados en raíz', 'Eliminados',
     'Eran copias de las imágenes en public/images/productos/. No referenciados por el código.'],
    ['error.txt', 'Eliminado',
     'Log de errores en tiempo de ejecución.'],
    ['php_errors.log', 'Eliminado',
     'Log de errores PHP.'],
    ['FitStock-API/tmp_test_insert.php', 'Eliminado',
     'Archivo de prueba temporal no perteneciente al código de producción.'],
    ['FitStock-API/estructura.txt', 'Eliminado',
     'Documentación de estructura desactualizada (referenciaba directorios que ya no existen).'],
    ['admin-sidebar.component.ts', 'Renombrado a admin-sidebar.ts',
     'Inconsistencia de nomenclatura: todos los demás componentes usan nombre plano (ej. login.ts), este usaba .component.ts.'],
]

for r in rows_data4:
    row = table4.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph('')

# ============================================================
# 3. ERRORES Y SOLUCIONES
# ============================================================
doc.add_heading('3. Errores Encontrados y Soluciones', level=1)

# Error 1
doc.add_heading('3.1 Error: "Call to undefined function mb_strtolower()"', level=2)
p = doc.add_paragraph()
p.add_run('Síntoma: ').bold = True
p.add_run('Al crear un material prestable, la API devolvía error 500. El log de PHP mostraba:\n')
p.add_run('"Call to undefined function mb_strtolower()" en Material.php línea 30.').italic = True

p2 = doc.add_paragraph()
p2.add_run('Causa: ').bold = True
p2.add_run('El método generarQr() usaba funciones mb_* (mb_strtolower, mb_strtoupper, mb_substr) que requieren la extensión PHP mbstring, no instalada en el servidor.')

p3 = doc.add_paragraph()
p3.add_run('Solución: ').bold = True
p3.add_run('Se reemplazaron las funciones multibyte por sus equivalentes estándar (strtolower, strtoupper, substr), dado que los nombres de material son ASCII y no necesitan soporte multibyte.')

p4 = doc.add_paragraph()
p4.add_run('Archivo modificado: ').bold = True
p4.add_run('FitStock-API/models/Material.php — línea 30.')

doc.add_paragraph('')

# Error 2
doc.add_heading('3.2 Error: Imágenes de productos no se mostraban', level=2)
p = doc.add_paragraph()
p.add_run('Síntoma: ').bold = True
p.add_run('Las imágenes de los productos aparecían como rotas (404) en las tarjetas.')

p2 = doc.add_paragraph()
p2.add_run('Causa: ').bold = True
p2.add_run('(1) Las imágenes estaban en la raíz del proyecto en lugar de public/images/productos/. '
           '(2) Los nombres de producto contienen espacios (ej. "Barrita proteica chocolate.jpg") pero las URL no estaban codificadas.')

p3 = doc.add_paragraph()
p3.add_run('Solución: ').bold = True
p3.add_run('Se movieron todas las imágenes a public/images/productos/ y se añadió encodeURIComponent() en el método getImagenUrl() para codificar los espacios.')

p4 = doc.add_paragraph()
p4.add_run('Archivo modificado: ').bold = True
p4.add_run('producto-list.ts — método getImagenUrl().')

doc.add_paragraph('')

# Error 3
doc.add_heading('3.3 Error: Precio devuelto como string en lugar de número', level=2)
p = doc.add_paragraph()
p.add_run('Síntoma: ').bold = True
p.add_run('El precio del producto se mostraba como string, causando problemas al calcular totales o al usar toFixed().')

p2 = doc.add_paragraph()
p2.add_run('Causa: ').bold = True
p2.add_run('PHP devolvía el valor DECIMAL de MySQL como string en JSON.')

p3 = doc.add_paragraph()
p3.add_run('Solución: ').bold = True
p3.add_run('Se añadió floatval() en la API de PHP al devolver el precio. En la plantilla Angular se usa (+p.precio).toFixed(2) como doble validación.')

p4 = doc.add_paragraph()
p4.add_run('Archivos modificados: ').bold = True
p4.add_run('API index.php, template producto-list.html.')

doc.add_paragraph('')

# Error 4
doc.add_heading('3.4 Warning: Presupuesto CSS excedido (producto-list.css)', level=2)
p = doc.add_paragraph()
p.add_run('Síntoma: ').bold = True
p.add_run('Al compilar Angular: "producto-list.css exceeded maximum budget. Budget 4.00 kB was not met by 1.55 kB with a total of 5.55 kB."')

p2 = doc.add_paragraph()
p2.add_run('Causa: ').bold = True
p2.add_run('El CSS rediseñado con las tarjetas, grid responsive y estilos del carrito supera el límite de 4 kB configurado en angular.json.')

p3 = doc.add_paragraph()
p3.add_run('Solución: ').bold = True
p3.add_run('No es un error bloqueante. Se puede aumentar el budget en angular.json o minimizar el CSS si es necesario.')

doc.add_paragraph('')

# ============================================================
# 4. ESTRUCTURA FINAL DEL PROYECTO
# ============================================================
doc.add_heading('4. Estructura Final del Proyecto', level=1)

doc.add_paragraph('FitStock-API/ (Backend PHP)')
items_api = [
    'api/index.php — Enrutador de la API REST',
    'conexion.php — Conexión a MySQL',
    'router.php — Router HTTP para el servidor de desarrollo',
    'config/database.sql — Esquema de base de datos',
    'models/ — Modelos: Acceso, Incidencia, Material, Prestamo, Producto, Usuario',
]
for item in items_api:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('FitStock-APP/ (Frontend Angular)')
items_app = [
    'src/app/components/ — 12 componentes (cada uno con .ts, .html, .css)',
    'src/app/services/ — 6 servicios (materiales, prestamos, productos, incidencias, usuario, mock-data)',
    'src/app/interfaces/ — interfaces compartidas',
    'public/images/productos/ — 8 imágenes de productos',
    'public/icono.jpg — Icono de la aplicación',
]
for item in items_app:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 5. ESTADO ACTUAL
# ============================================================
doc.add_heading('5. Estado Actual', level=1)

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Light Grid Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0].cells
hdr5[0].text = 'Componente'
hdr5[1].text = 'Estado'
hdr5[2].text = 'Observaciones'

rows_data5 = [
    ['Angular build', '✅ Correcto', 'Sin errores. Warning no crítico de budget CSS.'],
    ['PHP lint', '✅ Correcto', '9 archivos sin errores de sintaxis.'],
    ['Productos (Inventario)', '✅ Completo', 'Tarjetas con imágenes, carrito de compra, CRUD completo.'],
    ['Préstamos', '✅ Completo', 'Creación de material, préstamo con selector de usuario, edición y borrado de materiales.'],
    ['Incidencias', '✅ Completo', 'CRUD de incidencias.'],
    ['Usuarios', '✅ Completo', 'CRUD de usuarios, login, registro, perfil.'],
    ['Base de datos', '✅ Sincronizada', 'database.sql coincide con Docker.'],
    ['Limpieza de archivos', '✅ Completa', 'Eliminados duplicados, temporales y desactualizados.'],
]

for r in rows_data5:
    row = table5.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph('')

# ============================================================
# 6. PRÓXIMOS PASOS
# ============================================================
doc.add_heading('6. Próximos Pasos', level=1)
pasos = [
    'Subir imágenes faltantes: "Bebida energetica naranja", "Bebida energetica limon", "Creatina monohidrato 500g".',
    'Aumentar el budget CSS en angular.json o minimizar producto-list.css.',
    'Añadir .gitignore para dist/ y otros artefactos de build.',
    'Revisar si el modelo Acceso.php es necesario o eliminarlo.',
    'Pruebas de integración completa (flujo préstamo + devolución).',
]
for p in pasos:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph('')

# -- Guardar --
output_path = 'C:\\TFG\\tfg-FitStock\\Reporte_Cambios_Errores.docx'
doc.save(output_path)
print(f'Reporte generado en: {output_path}')
